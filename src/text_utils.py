from pdfminer.converter import PDFConverter
from pdfminer.high_level import extract_text as pdf_text
from docx import Document
from pdfminer.pdfdocument import PDFNoValidXRef, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFSyntaxError, PDFParser
from pathlib import Path
import unicodedata
import re
import logging
import io

# Silencing spam from pdfminer
logging.getLogger("pdfminer").setLevel(logging.CRITICAL)


class VectorCounter(PDFConverter):
    """
    A custom PDFMiner device that doesn't extract text,
    but counts how many vector drawing instructions it sees.
    """

    def __init__(self, rsrcmgr, outfp, codec='utf-8', pageno=1, laparams=None):
        super().__init__(rsrcmgr, outfp, codec=codec, pageno=pageno, laparams=laparams)
        self.vector_ops_count = 0

    # These methods are called by the interpreter when it encounters drawing commands
    def paint_path(self, gstate, stroke, fill, evenodd, path):
        # Every time a line, curve, or shape is drawn, this triggers
        self.vector_ops_count += 1

    # We ignore text commands to speed things up
    def render_image(self, name, stream): pass

    def render_string(self, PDFTextState , PDFTextSeq, PDFColorSpace,PDFGraphicState): pass


def is_complex_vector_file(path_str, threshold=500):
    """
    Analyzes the first page. If it finds more than 'threshold' vector paths,
    it assumes the file is a CAD drawing/Vector graphic.
    """
    try:
        with open(path_str, 'rb') as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            outfp = io.StringIO()

            # Use our custom counter device
            device = VectorCounter(rsrcmgr, outfp)
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            # Only process the first page
            for page in PDFPage.create_pages(doc):
                interpreter.process_page(page)
                break  # One page is enough to know

            count = device.vector_ops_count
            device.close()
            outfp.close()

            if count > threshold:
                print(f"[INFO] Detected Vector Drawing (paths: {count}): {path_str}")
                return True

            return False
    except Exception:
        # If we can't count vectors (e.g. file corrupt), assume it's unsafe.
        return True


def is_created_by_cad_software(path):
    try:
        with open(path, 'rb') as fp:
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            # Grabing metadata
            info = doc.info[0] if doc.info else {}

            def decode_meta(val):
                """Decodes meta information from bytes."""
                if isinstance(val, bytes):
                    return val.decode('utf-8', 'ignore').lower()
                return str(val).lower()

            producer = decode_meta(info.get('Producer', '')).lower()
            creator = decode_meta(info.get('Creator', '')).lower()

            # If we see these names, we ABORT before reading pages
            cad_signatures = ['autocad', 'bentley', 'microstation', 'revit', 'bluebeam', 'graphisoft']

            if any(sig in producer for sig in cad_signatures) or any(sig in creator for sig in cad_signatures):
                print(f"[INFO] Skipped CAD/Vector PDF: {path}")
                return True
            return
    # if it's broken this time skip the file
    except Exception:
        print(f"[WARN] Unable to open and skipping: {path}")
        return True


def extract_text(path : Path):
    """
    Identifies the file type by extension and attempts to extract its raw text content.

    This function acts as a wrapper for specific file handlers (PDFMiner, python-docx).
    It is designed to fail silently for corrupt files to allow the scanning process
    to continue uninterrupted.

    Args:
        path (str | Path): The file path to process.

    Returns:
        str | None: The extracted text string if successful. Returns None if:
            - The file extension is not supported (.txt, .pdf, .docx).
            - The file is corrupt or unreadable.
            - An exception occurs during parsing.
    """
    path_str = str(path)
    try:
        if path_str.endswith(".txt"):
            with open(path_str, "r", errors="ignore", encoding="utf-8") as f:
                return f.read()
        if path_str.endswith(".pdf"):
            try:
                # Peeking into .pdf file
                if is_created_by_cad_software(path_str):
                    return None
                # Opening file and counting vectors
                if is_complex_vector_file(path_str):
                    return None
                # finnaly there were no problems
                return pdf_text(path_str)
            except (PDFSyntaxError, PDFNoValidXRef, Exception):
                return None
        if path_str.endswith(".docx"):
            doc = Document(path_str)
            return "\n".join(p.text for p in doc.paragraphs)
        return None
    except Exception as e:
        print(f"[WARNING] Failed to extract {path_str}: {e}")
        return None


def text_clean(text_data: str) -> str:
    """
        Normalizes raw text data to facilitate fuzzy matching.

        This removes formatting differences that shouldn't count as "unique" content,
        such as capitalization, variable whitespace, or specific Unicode representations.

        Args:
            text_data (str): The raw input string.

        Returns:
            str: The normalized string. Returns an empty string if input is None or empty.

        Transformations:
            1. Unicode Normalization (NFKD): Decomposes characters (e.g., splits accented
               characters into base char + combining diacritic).
            2. Lowercasing: Converts all characters to lowercase.
            3. Whitespace Collapsing: Converts tabs, newlines, and multi-spaces into
               a single space character.
        """
    if not text_data:
        return ""
    text_data = unicodedata.normalize("NFKD", text_data)
    text_data = text_data.lower()
    text_data = re.sub(r"\s+", " ", text_data)
    return text_data.strip()